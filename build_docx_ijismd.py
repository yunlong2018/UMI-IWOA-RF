# -*- coding: utf-8 -*-
"""
Build an IJISMD (IGI Global)-style, single-column Word version of the manuscript.

Changes applied relative to the Markdown manuscript:
  - numbered citations [n] are converted to APA author-date in text;
  - the reference list is rewritten in APA 7th style (alphabetical);
  - generated figures (Fig. 1-5) are embedded;
  - tables are rendered with borders (adapt to the official IGI template if required).

Usage:
  python build_docx_ijismd.py
Output:
  ../Manuscript_IJISMD_format.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent
ROOT = BASE.parent
MD = ROOT / "Manuscript_University_Management_Informatization_IWOA-RF.md"
OUT_DOCX = ROOT / "Manuscript_IJISMD_format.docx"
FIG_DIR = BASE / "results" / "figures"
FIG_FILES = {
    "1": "fig1_workflow.png",
    "2": "fig2_architecture.png",
    "3": "fig3_confusion_matrix.png",
    "4": "fig4_metrics_comparison.png",
    "5": "fig5_roc_curve.png",
}

# APA author-date strings keyed by normalized citation ("5,6" -> "[5, 6]")
_APA = {
    "1": "(Ministry of Education of the People's Republic of China, 2018)",
    "2": "(UNESCO, 2021)",
    "3": "(Selwyn, 2016)",
    "4": "(Paulk et al., 1993)",
    "5": "(Janssen et al., 2012)",
    "6": "(Attard et al., 2015)",
    "7": "(Zachman, 1987)",
    "8": "(Pardo & Tayi, 2007)",
    "9": "(Zeeshan et al., 2022)",
    "10": "(Samala et al., 2025)",
    "11": "(DeLone & McLean, 2003)",
    "12": "(Hevner et al., 2004)",
    "13": "(March & Smith, 1995)",
    "14": "(Peffers et al., 2007)",
    "15": "(Lankhorst, 2013)",
    "16": "(Wang & Wang, 2025)",
    "17": "(Wang & Wang, 2026)",
    "18": "(Mirjalili & Lewis, 2016)",
    "19": "(Breiman, 2001)",
    "20": "(Cortes & Vapnik, 1995)",
    "21": "(Rumelhart et al., 1986)",
    "22": "(Friedman, 2001)",
    "23": "(Chen & Guestrin, 2016)",
    "24": "(Holland, 1992)",
    "25": "(Kennedy & Eberhart, 1995)",
    "26": "(Zuiderwijk & Janssen, 2014)",
    "27": "(Benbya et al., 2020)",
    "28": "(Venable et al., 2016)",
    "5,6": "(Attard et al., 2015; Janssen et al., 2012)",
    "7,15": "(Lankhorst, 2013; Zachman, 1987)",
    "9,10": "(Samala et al., 2025; Zeeshan et al., 2022)",
    "12,13": "(Hevner et al., 2004; March & Smith, 1995)",
    "16,17": "(Wang & Wang, 2025, 2026)",
}
_CITE_RE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")


def apa_cites(text):
    def _repl(m):
        key = ",".join(x.strip() for x in m.group(1).split(","))
        return _APA.get(key, m.group(0))
    return _CITE_RE.sub(_repl, text)


# APA 7th reference list (alphabetical). *...* marks italics.
REFS_APA = [
    "Attard, J., Orlandi, F., Scerri, S., & Auer, S. (2015). A systematic review of open "
    "government data initiatives. *Government Information Quarterly*, *32*(4), 399–418.",
    "Benbya, H., Nan, N., Tanriverdi, H., & Yoo, Y. (2020). Complexity and information "
    "systems research in the emerging digital world. *MIS Quarterly*, *44*(1), 1–17.",
    "Breiman, L. (2001). Random forests. *Machine Learning*, *45*(1), 5–32. "
    "https://doi.org/10.1023/A:1010933404324",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In "
    "*Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery "
    "and Data Mining* (pp. 785–794). https://doi.org/10.1145/2939672.2939785",
    "Cortes, C., & Vapnik, V. (1995). Support-vector networks. *Machine Learning*, *20*(3), "
    "273–297. https://doi.org/10.1007/BF00994018",
    "DeLone, W. H., & McLean, E. R. (2003). The DeLone and McLean model of information "
    "systems success: A ten-year update. *Journal of Management Information Systems*, "
    "*19*(4), 9–30. https://doi.org/10.1080/07421222.2003.11045748",
    "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. "
    "*The Annals of Statistics*, *29*(5), 1189–1232. https://doi.org/10.1214/aos/1013203451",
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information "
    "systems research. *MIS Quarterly*, *28*(1), 75–105. https://doi.org/10.2307/25148625",
    "Holland, J. H. (1992). *Adaptation in natural and artificial systems*. MIT Press.",
    "Janssen, M., Charalabidis, Y., & Zuiderwijk, A. (2012). Benefits, adoption barriers and "
    "myths of open data and open government. *Information Systems Management*, *29*(4), "
    "258–268.",
    "Kennedy, J., & Eberhart, R. (1995). Particle swarm optimization. In *Proceedings of "
    "ICNN'95 – International Conference on Neural Networks* (pp. 1942–1948). "
    "https://doi.org/10.1109/ICNN.1995.488968",
    "Lankhorst, M. (2013). *Enterprise architecture at work: Modelling, communication and "
    "analysis* (3rd ed.). Springer.",
    "March, S. T., & Smith, G. F. (1995). Design and natural science research on information "
    "technology. *Decision Support Systems*, *15*(4), 251–266. "
    "https://doi.org/10.1016/0167-9236(94)00041-O",
    "Ministry of Education of the People's Republic of China. (2018). *Education "
    "informatization 2.0 action plan* (Jiao Ji [2018] No. 6). Author.",
    "Mirjalili, S., & Lewis, A. (2016). The whale optimization algorithm. *Advances in "
    "Engineering Software*, *95*, 51–67. https://doi.org/10.1016/j.advengsoft.2016.01.008",
    "Pardo, T. A., & Tayi, G. K. (2007). Interorganizational information integration: A key "
    "enabler for digital government. *Government Information Quarterly*, *24*(4), 691–715.",
    "Paulk, M. C., Curtis, B., Chrissis, M. B., & Weber, C. V. (1993). *Capability maturity "
    "model for software, version 1.1* (CMU/SEI-93-TR-024). Software Engineering Institute, "
    "Carnegie Mellon University.",
    "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design "
    "science research methodology for information systems research. *Journal of Management "
    "Information Systems*, *24*(3), 45–77. https://doi.org/10.2753/MIS0742-1222240302",
    "Rumelhart, D. E., Hinton, G. E., & Williams, R. J. (1986). Learning representations by "
    "back-propagating errors. *Nature*, *323*, 533–536. https://doi.org/10.1038/323533a0",
    "Samala, A. D., Rawas, S., Rahmadika, S., & Indarta, Y. (2025). Virtual reality in "
    "education: Global trends, challenges, and impacts—game changer or passing trend? "
    "*Discover Education*, *4*, 229.",
    "Selwyn, N. (2016). *Education and technology: Key issues and debates* (2nd ed.). "
    "Bloomsbury. ISBN 978-1-4742-3591-4.",
    "UNESCO. (2021). *Reimagining our futures together: A new social contract for "
    "education*. UNESCO. https://doi.org/10.54675/ASRB4722",
    "Venable, J., Pries-Heje, J., & Baskerville, R. (2016). FEDS: A framework for "
    "evaluation in design science research. *European Journal of Information Systems*, "
    "*25*(1), 77–89. https://doi.org/10.1057/ejis.2015.7",
    "Wang, Y., & Wang, Z. (2025). Research on information sharing and tracking platform for "
    "educational management based on Internet of Things technology. *Discover Artificial "
    "Intelligence*, *5*, 320. https://doi.org/10.1007/s44163-025-00594-1",
    "Wang, Y., & Wang, Z. (2026). Application research and analysis of virtual reality "
    "technology in educational management training. *Discover Artificial Intelligence*, "
    "*6*, 139. https://doi.org/10.1007/s44163-025-00627-9",
    "Zachman, J. A. (1987). A framework for information systems architecture. *IBM Systems "
    "Journal*, *26*(3), 276–292.",
    "Zeeshan, K., Hämäläinen, T., & Neittaanmäki, P. (2022). Internet "
    "of things for sustainable smart education: An overview. *Sustainability*, *14*(7), "
    "4293. https://doi.org/10.3390/su14074293",
    "Zuiderwijk, A., & Janssen, M. (2014). Open data policies, their implementation and "
    "impact: A framework for comparison. *Government Information Quarterly*, *31*(1), "
    "17–29.",
]


def add_bold_runs(p, text):
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.bold = (idx % 2 == 1)


def add_italic_runs(p, text):
    parts = text.split("*")
    for idx, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.italic = (idx % 2 == 1)


def add_table(doc, rows):
    n_rows = len(rows)
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = tbl.cell(ri, ci)
            p = cell.paragraphs[0]
            bold = cell_text.startswith("**") and cell_text.endswith("**")
            txt = cell_text[2:-2] if bold else cell_text
            r = p.add_run(apa_cites(txt))
            r.bold = bold or ri == 0
            r.font.size = Pt(9.5)


def main():
    if not MD.exists():
        raise SystemExit(f"Manuscript not found: {MD}")
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for hs, sz, italic in [("Heading 1", 14, False), ("Heading 2", 12, False),
                           ("Heading 3", 11, True)]:
        st = doc.styles[hs]
        st.font.name = "Times New Roman"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.italic = italic
        st.font.color.rgb = RGBColor(0, 0, 0)

    in_code = False
    code_buf = []
    refs_mode = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if in_code:
            if line.strip().startswith("```"):
                for c in code_buf:
                    p = doc.add_paragraph()
                    r = p.add_run(c)
                    r.font.name = "Consolas"
                    r.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.3)
                in_code = False
            else:
                code_buf.append(line)
            i += 1
            continue
        if line.strip().startswith("```"):
            in_code = True
            code_buf = []
            i += 1
            continue
        s = line.strip()
        if not s:
            i += 1
            continue
        if refs_mode:
            i += 1
            continue
        if s.startswith("# "):
            p = doc.add_heading(s[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if s.startswith("## "):
            title = s[3:].strip()
            doc.add_heading(title, level=1)
            if title == "References":
                refs_mode = True
                for ref in REFS_APA:
                    p = doc.add_paragraph()
                    add_italic_runs(p, ref)
                    pf = p.paragraph_format
                    pf.left_indent = Inches(0.5)
                    pf.first_line_indent = Inches(-0.5)
            i += 1
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=2)
            i += 1
            continue
        if s.startswith("**Table"):
            m = re.match(r"^\*\*Table\s*(\d+)\*\*\s*(.*)$", s)
            if m:
                p = doc.add_paragraph()
                r1 = p.add_run(f"Table {m.group(1)}. ")
                r1.bold = True
                r2 = p.add_run(m.group(2))
                r2.italic = True
                i += 1
                while i < len(lines) and lines[i].strip() == "":
                    i += 1  # allow a blank line between the caption and the table body
                rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    if not all(set(c) <= {"-"} for c in cells if c):
                        rows.append(cells)
                    i += 1
                if rows:
                    add_table(doc, rows)
                continue
        if s.startswith("**Fig."):
            m = re.match(r"^\*\*Fig\.\s*(\d+)\*\*\s*(.*)$", s)
            if m:
                fname = FIG_FILES.get(m.group(1))
                if fname and (FIG_DIR / fname).exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(FIG_DIR / fname), width=Inches(5.5))
                p = doc.add_paragraph()
                r1 = p.add_run(f"Figure {m.group(1)}. ")
                r1.bold = True
                r2 = p.add_run(apa_cites(m.group(2)))
                r2.italic = True
                i += 1
                continue
        if s.startswith("!["):
            i += 1  # markdown image embed: the docx inserts figures at caption lines instead
            continue
        if s.startswith("|"):
            i += 1
            continue
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_bold_runs(p, apa_cites(s[2:]))
            i += 1
            continue
        if re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(s)
            i += 1
            continue
        if s.startswith("**["):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(s.replace("**", ""))
            r.bold = True
            i += 1
            continue
        if s.startswith("¹") or s.startswith("²") or s.startswith("*Correspondence"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(s.lstrip("*"))
            i += 1
            continue
        if re.match(r"^[^#|\-*\d]", s) and re.search(r"\(\d+\)$", s):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(s)
            i += 1
            continue
        p = doc.add_paragraph()
        add_bold_runs(p, apa_cites(s))
        i += 1

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
